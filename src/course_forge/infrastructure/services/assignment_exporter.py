import os
import subprocess
import mistune
import yaml
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
import re

try:
    from pygments import highlight
    from pygments.lexers import get_lexer_by_name, TextLexer
    from pygments.token import Token
    HAS_PYGMENTS = True
except ImportError:
    HAS_PYGMENTS = False

# Default location of the bundled template config (same package directory)
_DEFAULT_TEMPLATE_CONFIG = Path(__file__).parent.parent.parent / "templates" / "exam_template.yaml"

# Callout type -> (label, border hex color)
CALLOUT_STYLES = {
    "important": ("⚠ Importante",  "E57C00"),
    "warning":   ("⚠ Atenção",     "D32F2F"),
    "danger":    ("✖ Perigo",      "B71C1C"),
    "note":      ("ℹ Nota",        "1976D2"),
    "tip":       ("✔ Dica",        "2E7D32"),
    "info":      ("ℹ Info",        "0288D1"),
}

# Pygments token -> (bold, italic, hex_rgb or None)
_PYGMENTS_TOKEN_STYLE = {
    Token.Keyword:            (True,  False, "0000CC"),
    Token.Keyword.Type:       (True,  False, "0000CC"),
    Token.Name.Builtin:       (True,  False, "0000CC"),
    Token.Name.Function:      (False, False, "0000AA"),
    Token.Name.Class:         (True,  False, "006600"),
    Token.Comment:            (False, True,  "888888"),
    Token.Comment.Single:     (False, True,  "888888"),
    Token.Comment.Multiline:  (False, True,  "888888"),
    Token.Literal.String:     (False, False, "CC0000"),
    Token.Literal.String.Doc: (False, True,  "880000"),
    Token.Literal.Number:     (False, False, "AA00AA"),
    Token.Operator:           (False, False, "555555"),
    Token.Operator.Word:      (True,  False, "0000CC"),
    Token.Punctuation:        (False, False, "333333"),
    Token.Generic.Heading:    (True,  False, "000080"),
}

def _hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

class AssignmentExporter:
    def __init__(self, template_path=None, template_config=None):
        """Create the exporter.

        Parameters
        ----------
        template_path:
            Optional path to an existing .docx file to use as base template.
            When None (default) the document is built programmatically from
            the bundled YAML configuration.
        template_config:
            Optional path to a YAML config file. Defaults to the bundled
            ``exam_template.yaml`` inside the package.
        """
        self.template_path = template_path  # kept for backwards-compat (may be None)
        self.total_points = 0.0
        self.out_dir = ""

        # Load template config from YAML
        config_path = Path(template_config) if template_config else _DEFAULT_TEMPLATE_CONFIG
        with open(config_path, "r", encoding="utf-8") as f:
            self._tpl = yaml.safe_load(f)

    # ------------------------------------------------------------------
    # Template document builder
    # ------------------------------------------------------------------

    def _build_base_document(self, is_exam: bool, metadata: dict = None) -> Document:
        """Build the base Word document from the YAML config, no external file needed."""
        tpl = self._tpl
        doc = Document()
        if metadata is None: metadata = {}

        # --- Page setup ---
        page_cfg = tpl.get("page", {})
        section = doc.sections[0]
        if page_cfg.get("width"):
            section.page_width = Emu(page_cfg["width"])
        if page_cfg.get("height"):
            section.page_height = Emu(page_cfg["height"])
        margins = page_cfg.get("margins", {})
        if margins.get("top") is not None:
            section.top_margin = Emu(margins["top"])
        if margins.get("bottom") is not None:
            section.bottom_margin = Emu(margins["bottom"])
        if margins.get("left") is not None:
            section.left_margin = Emu(margins["left"])
        if margins.get("right") is not None:
            section.right_margin = Emu(margins["right"])
        if page_cfg.get("footer_distance") is not None:
            section.footer_distance = Emu(page_cfg["footer_distance"])
        if page_cfg.get("header_distance") is not None:
            section.header_distance = Emu(page_cfg["header_distance"])

        # --- FA-DIF-005 code in top-right ---
        hdr_cfg = tpl.get("header", {})
        form_code = hdr_cfg.get("form_code", "")
        if form_code:
            code_para = doc.add_paragraph()
            code_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            code_run = code_para.add_run(form_code)
            code_run.font.size = Pt(hdr_cfg.get("form_code_size_pt", 12))
            # Empty paragraph after code
            doc.add_paragraph()

        # --- Header table (identification) ---
        htbl_cfg = tpl.get("header_table", {})
        font_pt = htbl_cfg.get("font_size_pt", 10)
        # Build table based on type
        num_rows = 3 if is_exam else 2
        htable = doc.add_table(rows=num_rows, cols=4)
        try:
            htable.style = htbl_cfg.get("style", "Table Grid")
        except Exception:
            pass

        # Disable autofit to force specific column widths
        htable.autofit = False
        tblPr = htable._tbl.tblPr
        
        # Set exact table layout
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)
        
        # Configure table width to exactly 10240 dxa (approx 16 cm)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '10240')
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

        # Force the column grid (tblGrid) to obey our layout
        # Col 0 (Logo): 2400 dxa
        # Total = 10240 dxa
        if is_exam:
            dxas = ['2400', '5280', '1280', '1280']
        else:
            # Shift space to Entrega (Col 2: 2720 dxa) taking from Titles (Col 1: 3840 dxa)
            dxas = ['2400', '3840', '2720', '1280']
            
        tblGrid = OxmlElement('w:tblGrid')
        for w_dxa in dxas:
            gridCol = OxmlElement('w:gridCol')
            gridCol.set(qn('w:w'), w_dxa)
            tblGrid.append(gridCol)
            
        # Remove any existing tblGrid that python-docx might have added and append ours
        existing_grid = htable._tbl.find(qn('w:tblGrid'))
        if existing_grid is not None:
            htable._tbl.remove(existing_grid)
        htable._tbl.append(tblGrid)

        # Row 0: cell[0] = logo placeholder, cells[1-3] merged = Curso/Disciplina
        r0 = htable.rows[0]
        
        # Exact values in EMUs corresponding to the dxa above (dxa * 635)
        c0_w = int(dxas[0]) * 635
        c1_w = int(dxas[1]) * 635
        c2_w = int(dxas[2]) * 635
        c3_w = int(dxas[3]) * 635
        
        # We must set widths on standard row objects so python-docx renders strictly
        r0.cells[0].width = Emu(c0_w)
        r0.cells[1].width = Emu(c1_w)
        r0.cells[2].width = Emu(c2_w)
        r0.cells[3].width = Emu(c3_w)

        # Set custom margins for logo cell: 0 top/bottom, 108 dxa left/right (standard word padding)
        tcPr_logo = r0.cells[0]._tc.get_or_add_tcPr()
        tcMar_logo = OxmlElement('w:tcMar')
        for m in ['top', 'bottom', 'left', 'right']:
            node = OxmlElement(f'w:{m}')
            if m in ['top', 'bottom']:
                node.set(qn('w:w'), '0')
            else:
                node.set(qn('w:w'), '108')
            node.set(qn('w:type'), 'dxa')
            tcMar_logo.append(node)
        tcPr_logo.append(tcMar_logo)

        # Insert logo if specified
        logo_path = tpl.get("institution", {}).get("logo")
        if logo_path:
            logo_abs = Path(_DEFAULT_TEMPLATE_CONFIG).parent.parent / logo_path
            p_logo = r0.cells[0].paragraphs[0]
            p_logo.clear()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if logo_abs.exists():
                r_logo = p_logo.add_run()
                # Diminuídas para não encostar na borda esquerda/direita (1.25M emu aprox)
                r_logo.add_picture(str(logo_abs), width=Emu(1250000), height=Emu(617220))
            else:
                p_logo.text = "[LOGO]"
        else:
            r0.cells[0].text = ""

        # Rest of row widths must be explicit for word to render correctly when merged
        r1 = htable.rows[1]
        
        # Merge columns 1-3 for Curso/Disciplina
        merged_01 = r0.cells[1].merge(r0.cells[3])
        merged_01.width = Emu(c1_w + c2_w + c3_w)
        p_curso = merged_01.paragraphs[0]
        p_curso.clear()
        rr = p_curso.add_run("Curso:")
        rr.bold = True; rr.font.size = Pt(font_pt)
        merged_01.add_paragraph()  # blank line between Curso and Disciplina
        p_disc = merged_01.add_paragraph()
        rd = p_disc.add_run("Disciplina:")
        rd.bold = True; rd.font.size = Pt(font_pt)

        # Define date label
        date_label = htbl_cfg.get("date_label", "Data:")

        if is_exam:
            r2 = htable.rows[2]
            # Row 1: cells[0-1] merged = Professor(a), cells[2-3] merged = Data
            merged_prof = r1.cells[0].merge(r1.cells[1])
            merged_prof.width = Emu(c0_w + c1_w)
            r1.height = Pt(30)
            p_prof = merged_prof.paragraphs[0]
            p_prof.clear()
            rp = p_prof.add_run("Professor(a):")
            rp.bold = True; rp.font.size = Pt(font_pt)
            
            merged_data = r1.cells[2].merge(r1.cells[3])
            merged_data.width = Emu(c2_w + c3_w)
            p_data = merged_data.paragraphs[0]
            p_data.clear()
            rd2 = p_data.add_run(date_label)
            rd2.bold = True; rd2.font.size = Pt(font_pt)

            # Row 2: cells[0-1] merged = Nome aluno, cell[2] = Valor, cell[3] = Nota
            merged_nome = r2.cells[0].merge(r2.cells[1])
            merged_nome.width = Emu(c0_w + c1_w)
            r2.height = Pt(30)
            p_nome = merged_nome.paragraphs[0]
            p_nome.clear()
            rn = p_nome.add_run("Nome do(a) aluno(a):")
            rn.bold = True; rn.font.size = Pt(font_pt)
            
            p_val = r2.cells[2].paragraphs[0]
            r2.cells[2].width = Emu(c2_w)
            p_val.clear()
            rv = p_val.add_run("Valor:")
            rv.bold = True; rv.font.size = Pt(font_pt)
            
            p_nota = r2.cells[3].paragraphs[0]
            r2.cells[3].width = Emu(c3_w)
            p_nota.clear()
            rnt = p_nota.add_run("Nota:")
            rnt.bold = True; rnt.font.size = Pt(font_pt)
        else:
            # ASSIGNMENT mode: 
            # Row 1: cells[0-1] merged = Professor, cell[2] = Data, cell[3] = Valor
            merged_prof = r1.cells[0].merge(r1.cells[1])
            merged_prof.width = Emu(c0_w + c1_w)
            r1.height = Pt(30)
            p_prof = merged_prof.paragraphs[0]
            p_prof.clear()
            rp = p_prof.add_run("Professor(a):")
            rp.bold = True; rp.font.size = Pt(font_pt)
            
            r1.cells[2].width = Emu(c2_w)
            p_data = r1.cells[2].paragraphs[0]
            p_data.clear()
            
            due_date = metadata.get("due_date")
            if due_date:
                rd2 = p_data.add_run(f"Entrega: {due_date}")
            else:
                rd2 = p_data.add_run("Entrega:")
            
            rd2.bold = True; rd2.font.size = Pt(font_pt)
            
            r1.cells[3].width = Emu(c3_w)
            p_val = r1.cells[3].paragraphs[0]
            p_val.clear()
            rv = p_val.add_run("Valor:")
            rv.bold = True; rv.font.size = Pt(font_pt)

        # Vertically center all cells in the header table
        # Para tabelas mescladas, precisamos garantir que tcPr existe
        for r in htable.rows:
            for c in r.cells:
                tcPr = c._tc.get_or_add_tcPr()
                
                # Remover vAlign se já existir e adicionar um novo
                for existing_vAlign in tcPr.findall(qn('w:vAlign')):
                    tcPr.remove(existing_vAlign)
                
                vAlign = OxmlElement('w:vAlign')
                vAlign.set(qn('w:val'), 'center')
                tcPr.append(vAlign)

                # Resetar espaçamentos do parágrafo pra não desalinhar
                for p in c.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)

        # Spacing after header table
        space_after_pt = hdr_cfg.get("table_spacing_after", 0)
        if space_after_pt > 0:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_after = Pt(space_after_pt)
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.line_spacing = 1.0

        # --- Instructions table (only in exam mode) ---
        instr_cfg = tpl.get("instructions", {})
        if instr_cfg.get("enabled", True) and is_exam:
            instr_table = doc.add_table(rows=1, cols=1)
            instr_table.style = "Normal Table"
            
            # Add borders and margins to the table
            tblPr = instr_table._tbl.tblPr
            
            # Set table width to 10245 dxa
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '10245')
            tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)
            
            # Align center
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            tblPr.append(jc)

            # Borders
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '8')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tblBorders.append(border)
            tblPr.append(tblBorders)
            
            # Cell margins (100 dxa = very small margin)
            tblCellMar = OxmlElement('w:tblCellMar')
            for m in ['top', 'left', 'bottom', 'right']:
                mar = OxmlElement(f'w:{m}')
                mar.set(qn('w:w'), '100')
                mar.set(qn('w:type'), 'dxa')
                tblCellMar.append(mar)
            tblPr.append(tblCellMar)

            cell = instr_table.rows[0].cells[0]

            # Title
            title_p = cell.paragraphs[0]
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_p.paragraph_format.space_after = Pt(6)
            title_p.paragraph_format.space_before = Pt(4)
            tr = title_p.add_run(instr_cfg.get("title", "Instruções"))
            tr.bold = True
            tr.font.size = Pt(instr_cfg.get("font_size_pt", 9))

            # Instruction items
            font_name = instr_cfg.get("font_name", "Calibri")
            for item in instr_cfg.get("items", []):
                ip = cell.add_paragraph()
                ip.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                ip.style = "Normal" # Remove List Paragraph style to avoid huge indents
                ip.paragraph_format.space_after = Pt(2)
                ip.paragraph_format.space_before = Pt(2)
                ir = ip.add_run(item["text"])
                ir.font.name = font_name
                ir.font.size = Pt(instr_cfg.get("font_size_pt", 9))

            # Add spacing after instructions table
            post_instr = doc.add_paragraph()
            post_instr.paragraph_format.space_before = Pt(12)
            post_instr.paragraph_format.space_after = Pt(12)

        # --- Answer grid table (only in exam mode) ---
        ag_cfg = tpl.get("answer_grid", {})
        if ag_cfg.get("enabled", True) and is_exam:
            n = ag_cfg.get("num_questions", 10)
            ag_font_pt = ag_cfg.get("font_size_pt", 10)
            ag_table = doc.add_table(rows=3, cols=n)
            ag_table.style = "Normal Table"

            # Row 0: all cells merged into one header spanning all columns
            header_cell = ag_table.rows[0].cells[0].merge(ag_table.rows[0].cells[n - 1])
            for i, text in enumerate([
                ag_cfg.get("title", "QUADRO-RESPOSTA"),
                ag_cfg.get("subtitle", ""),
                ag_cfg.get("instruction", ""),
            ]):
                p = header_cell.paragraphs[0] if i == 0 else header_cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(text)
                r.bold = True
                r.font.size = Pt(ag_font_pt)

            # Row 1: question numbers 01..n
            for ci in range(n):
                p = ag_table.rows[1].cells[ci].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f"{ci + 1:02d}")
                r.bold = True
                r.font.size = Pt(ag_font_pt)

            # Row 2: empty answer cells
            for ci in range(n):
                p = ag_table.rows[2].cells[ci].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        return doc

    def _hex_to_rgb(self, h):
        h = h.lstrip("#")
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    def _add_left_border(self, paragraph, hex_color):
        """Add a colored left border to a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')   # 3pt line
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), hex_color)
        pBdr.append(left)
        pPr.append(pBdr)

    def _add_shading(self, paragraph, hex_color):
        """Add a light background shading."""
        # Lighten the border colour by mixing with white at 85%
        r, g, b = self._hex_to_rgb(hex_color)
        lr = int(r * 0.15 + 255 * 0.85)
        lg = int(g * 0.15 + 255 * 0.85)
        lb = int(b * 0.15 + 255 * 0.85)
        fill = f"{lr:02X}{lg:02X}{lb:02X}"
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        pPr.append(shd)

    def _delete_paragraph(self, paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    def _add_runs(self, paragraph, children):
        for child in children:
            child_type = child.get("type")
            if child_type == "text":
                text = child.get("raw", child.get("text", ""))
                run = paragraph.add_run(text)
                run.font.name = "Arial"
                run.font.size = Pt(10)
            elif child_type == "softbreak":
                # Handled by chunking now, but fallback to space
                run = paragraph.add_run(" ")
                run.font.name = "Arial"
                run.font.size = Pt(10)
            elif child_type == "strong":
                full_text = "".join(grand.get("raw", grand.get("text", "")) for grand in child.get("children", []))
                
                # Look for explicit points: "**1. [1.5] Avalie...**" or "**1. (2,0 pontos) ...**"
                match = re.search(r'^(\d+)\.\s*(?:\[|\()(\d+(?:[.,]\d+)?)(?:\s*pontos?)?(?:\]|\))\s*(.*)', full_text, flags=re.IGNORECASE)
                if match:
                    num = match.group(1)
                    pts = match.group(2).replace('.', ',')
                    rest = match.group(3)
                    try:
                        self.total_points += float(pts.replace(',', '.'))
                    except ValueError:
                        pass
                else:
                    match = re.search(r'^(\d+)\.\s*(.*)', full_text)
                    if match:
                        num = match.group(1)
                        pts = "1,0"
                        rest = match.group(2)
                        self.total_points += 1.0
                    else:
                        num = None
                        
                if num is not None:
                    try:
                        val = float(pts.replace(',', '.'))
                    except ValueError:
                        val = 1.0
                    label_ponto = "ponto" if val == 1.0 else "pontos"
                    run = paragraph.add_run(f"Questão {num:0>2} ({pts} {label_ponto}): ")
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    run.bold = True
                    if rest:
                        run = paragraph.add_run(rest)
                        run.font.name = "Arial"
                        run.font.size = Pt(10)
                        run.bold = True
                else:
                    run = paragraph.add_run(full_text)
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    run.bold = True
            elif child_type == "emphasis":
                for grand in child.get("children", []):
                    run = paragraph.add_run(grand.get("raw", grand.get("text", "")))
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                    run.italic = True
            elif child_type == "codespan":
                # codespan is a leaf: raw holds the code text directly
                raw_code = child.get("raw", child.get("text", ""))
                if not raw_code and child.get("children"):
                    raw_code = "".join(c.get("raw", c.get("text", "")) for c in child["children"])
                run = paragraph.add_run(raw_code)
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            elif child_type == "link":
                self._add_runs(paragraph, child.get("children", []))
            elif child_type == "image":
                src = child.get("dest", child.get("src", ""))
                # Remove query params (like filename?raw=true)
                src = src.split('?')[0]
                # Embed image
                src_abs = os.path.join(self.out_dir, src)
                if src_abs.endswith('.svg') and os.path.exists(src_abs):
                    png_path = src_abs + ".png"
                    try:
                        subprocess.run(["convert", src_abs, png_path], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                        src_abs = png_path
                    except Exception as e:
                        print("Failed to convert SVG to PNG:", e)
                
                if os.path.exists(src_abs):
                    run = paragraph.add_run()
                    try:
                        # Add picture with reasonable width
                        run.add_picture(src_abs, width=Inches(5.0))
                    except Exception as e:
                        print(f"Failed to add picture {src_abs}:", e)
                else:
                    paragraph.add_run("[Image: " + child.get("alt", "") + "]")
            elif "children" in child:
                self._add_runs(paragraph, child["children"])

    def _render_node(self, node, doc, list_level=0, callout_style=None):
        ntype = node.get("type")
        
        if ntype == "heading":
            level = node.get("attrs", {}).get("level", 1)
            p = doc.add_paragraph()
            p.style = "Normal"
            if callout_style:
                self._add_left_border(p, callout_style)
                self._add_shading(p, callout_style)
            self._add_runs(p, node.get("children", []))
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(16 - level)

        elif ntype == "paragraph":
            chunks = []
            current_chunk = []
            for child in node.get("children", []):
                if child.get("type") == "softbreak":
                    chunks.append(current_chunk)
                    current_chunk = []
                else:
                    current_chunk.append(child)
            if current_chunk:
                chunks.append(current_chunk)
                
            for chunk in chunks:
                if not chunk: continue
                p = doc.add_paragraph()
                try:
                    p.style = "List Paragraph" if list_level > 0 else "Normal"
                except KeyError:
                    p.style = "Normal"
                
                # Check if this paragraph contains options to apply right styles
                is_option = False
                is_question = False
                for child in chunk:
                    if child.get("type") == "text":
                        text = child.get("raw", child.get("text", ""))
                        if re.match(r'^[a-e]\)', text.strip()):
                            is_option = True
                    elif child.get("type") == "strong":
                        full_text = "".join(grand.get("raw", grand.get("text", "")) for grand in child.get("children", []))
                        if re.search(r'^(\d+)\.\s+', full_text):
                            is_question = True
                            
                if is_option:
                    try: p.style = "Body Text" 
                    except: pass
                    p.paragraph_format.left_indent = Pt(14.2)  # around 0.5cm
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                else:
                    p.paragraph_format.space_after = Pt(6)
                    if is_question:
                        p.paragraph_format.space_before = Pt(18)
                        p.paragraph_format.keep_with_next = True
                    else:
                        p.paragraph_format.space_before = Pt(0)
    
                self._add_runs(p, chunk)

        elif ntype == "block_quote":
            callout_type = None
            title_nodes = []
            body_from_first_p = []
            
            children = node.get("children", [])
            if children and children[0].get("type") == "paragraph":
                p1_children = children[0].get("children", [])
                
                # 1. Identify if it's a callout and find marker end
                full_text_until_bracket = ""
                marker_node_idx = -1
                for i, c in enumerate(p1_children):
                    if c.get("type") == "text":
                        txt = c.get("raw", c.get("text", ""))
                        full_text_until_bracket += txt
                        if "]" in full_text_until_bracket:
                            marker_node_idx = i
                            break
                    elif c.get("type") == "softbreak": break
                    else: break
                
                marker_match = re.match(r'^\[!(\w+)\]', full_text_until_bracket)
                if marker_match:
                    callout_type = marker_match.group(1).lower()
                    
                    # 2. Split first paragraph into title line and body lines
                    softbreak_idx = -1
                    for i, c in enumerate(p1_children):
                        if c.get("type") == "softbreak":
                            softbreak_idx = i
                            break
                    
                    raw_title_nodes = p1_children[:softbreak_idx] if softbreak_idx != -1 else p1_children
                    body_from_first_p = p1_children[softbreak_idx+1:] if softbreak_idx != -1 else []
                    
                    # 3. Create stripped title nodes (copy to avoid mutating original AST)
                    # We need to remove "[!type] " from the beginning of the text nodes
                    import copy
                    title_nodes = copy.deepcopy(raw_title_nodes)
                    
                    marker_full_str = marker_match.group(0) # e.g. "[!warning]"
                    to_strip = len(marker_full_str)
                    
                    # Strip from the start of title_nodes
                    curr_idx = 0
                    while to_strip > 0 and curr_idx < len(title_nodes):
                        node_to_edit = title_nodes[curr_idx]
                        if node_to_edit.get("type") == "text":
                            node_text = node_to_edit.get("raw", node_to_edit.get("text", ""))
                            strip_now = min(len(node_text), to_strip)
                            new_text = node_text[strip_now:]
                            node_to_edit["raw"] = new_text
                            if "text" in node_to_edit: node_to_edit["text"] = new_text
                            to_strip -= strip_now
                        curr_idx += 1
                    
                    # If after marker there's a space, strip it too
                    if title_nodes and title_nodes[0].get("type") == "text":
                        t = title_nodes[0].get("raw", "")
                        title_nodes[0]["raw"] = t.lstrip()
                        if "text" in title_nodes[0]: title_nodes[0]["text"] = t.lstrip()

            label_info = CALLOUT_STYLES.get(callout_type)
            if label_info:
                label, border_hex = label_info
                # Title paragraph
                title_p = doc.add_paragraph()
                title_p.style = "Normal"
                title_p.paragraph_format.space_after = Pt(2)
                title_p.paragraph_format.space_before = Pt(6)
                title_p.paragraph_format.left_indent = Pt(12)
                self._add_left_border(title_p, border_hex)
                self._add_shading(title_p, border_hex)
                
                r = title_p.add_run(f"{label}")
                r.bold = True
                r.font.name = "Arial"
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(*self._hex_to_rgb(border_hex))
                
                if title_nodes:
                    title_p.add_run(" — ").bold = True
                    # Use _add_runs but we MUST ensure they are bold and correct color
                    start_runs_idx = len(title_p.runs)
                    self._add_runs(title_p, title_nodes)
                    for r in title_p.runs[start_runs_idx:]:
                        r.bold = True
                        r.font.name = "Arial"
                        r.font.size = Pt(10)
                        r.font.color.rgb = RGBColor(*self._hex_to_rgb(border_hex))

                # Body items
                if body_from_first_p:
                    bp = doc.add_paragraph()
                    bp.style = "Normal"
                    bp.paragraph_format.left_indent = Pt(12)
                    self._add_left_border(bp, border_hex)
                    self._add_shading(bp, border_hex)
                    self._add_runs(bp, body_from_first_p)
                
                for child in children[1:]:
                    self._render_node(child, doc, list_level, callout_style=border_hex)
                
                # Closing spacer
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(4)
                sp.paragraph_format.space_before = Pt(0)
            else:
                # Plain blockquote or failed callout detection
                border_hex = "999999"
                for child in children:
                    self._render_node(child, doc, list_level, callout_style=border_hex)

        elif ntype == "list":
            # For lists we map options slightly differently if they're a,b,c etc.
            is_ordered = node.get("attrs", {}).get("ordered", False)
            for idx, child in enumerate(node.get("children", [])):
                if child.get("type") == "list_item":
                    p = doc.add_paragraph()
                    if callout_style:
                        self._add_left_border(p, callout_style)
                        self._add_shading(p, callout_style)
                    
                    try:
                        # ...
                        p.style = "List Bullet" if not is_ordered else "List Number"
                    except KeyError:
                        p.style = "Normal"
                        p.paragraph_format.left_indent = Pt(list_level * 20 + 20 + (12 if callout_style else 0))
                        p.add_run("• ")
                    for grand in child.get("children", []):
                        if grand.get("type") == "block_text":
                            self._add_runs(p, grand.get("children", []))
                        else:
                            self._render_node(grand, doc, list_level + 1, callout_style=callout_style)

        elif ntype == "block_code":
            lang = node.get("attrs", {}).get("info", "") or ""
            code = node.get("raw", node.get("text", "")).rstrip("\n")

            def _shaded_para():
                cp = doc.add_paragraph()
                if callout_style:
                    self._add_left_border(cp, callout_style)
                    self._add_shading(cp, callout_style)
                
                cp.style = "Normal"
                cp.paragraph_format.space_after = Pt(0)
                cp.paragraph_format.space_before = Pt(0)
                # If inside callout, add more indent
                base_indent = 12 if callout_style else 0
                cp.paragraph_format.left_indent = Pt(base_indent + 8)
                
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F0F0F0')
                cp._p.get_or_add_pPr().append(shd)
                
                bdr = OxmlElement('w:pBdr')
                for side in ('top', 'left', 'bottom', 'right'):
                    el = OxmlElement(f'w:{side}')
                    el.set(qn('w:val'), 'single')
                    el.set(qn('w:sz'), '4')
                    el.set(qn('w:space'), '1')
                    el.set(qn('w:color'), 'CCCCCC')
                    bdr.append(el)
                cp._p.get_or_add_pPr().append(bdr)
                return cp

            if HAS_PYGMENTS and lang:
                try:
                    lexer = get_lexer_by_name(lang, stripall=True)
                except Exception:
                    lexer = None # Ensure lexer is None on exception
                if lexer:
                    # Tokenize and split into lines properly
                    lines = [[]]
                    for ttype, value in lexer.get_tokens(code):
                        parts = value.split('\n')
                        for i, part in enumerate(parts):
                            if part:
                                lines[-1].append((ttype, part))
                            if i < len(parts) - 1:
                                lines.append([])
                    
                    # Remove trailing empty line if it resulted from a final \n
                    if lines and not lines[-1]:
                        lines.pop()

                    for i, line_tokens in enumerate(lines):
                        lp = _shaded_para()
                        if i == len(lines) - 1:
                            lp.paragraph_format.space_after = Pt(6)
                        for ttype, val in line_tokens:
                            run = lp.add_run(val)
                            run.font.name = "Courier New"
                            run.font.size = Pt(9)
                            # Style resolution
                            style = None
                            tt = ttype
                            while tt:
                                style = _PYGMENTS_TOKEN_STYLE.get(tt)
                                if style: break
                                tt = tt.parent if hasattr(tt, 'parent') else None
                            if style:
                                b, it, h = style
                                run.bold, run.italic = b, it
                                if h: run.font.color.rgb = RGBColor(*self._hex_to_rgb(h))
                    return
            return

            # Fallback: plain monospace, one paragraph per line
            last_lp = None
            for line in code.split('\n'):
                last_lp = _shaded_para()
                r = last_lp.add_run(line)
                r.font.name = "Courier New"
                r.font.size = Pt(9)
            if last_lp is not None:
                last_lp.paragraph_format.space_after = Pt(6)

        elif ntype == "thematic_break":
            # Just ignore thematic break in assignments
            pass

        elif ntype == "table":
            children = node.get("children", [])
            head = next((c for c in children if c.get("type") == "table_head"), None)
            body = next((c for c in children if c.get("type") == "table_body"), None)

            head_rows = head.get("children", []) if head else []
            body_rows_raw = body.get("children", []) if body else []

            # body children may be table_row nodes directly or wrapped
            body_rows = []
            for item in body_rows_raw:
                if item.get("type") == "table_row":
                    body_rows.append(item.get("children", []))  # list of table_cell
                else:
                    body_rows.append([item])  # fallback

            header_cells_rows = []
            for hr in head_rows:
                if hr.get("type") in ("table_row", "table_head"):
                    header_cells_rows.append(hr.get("children", []))
                else:
                    # mistune wraps the head differently: head node IS the single row of cells
                    header_cells_rows.append(head_rows)
                    break

            all_row_cells = header_cells_rows + body_rows
            if not all_row_cells:
                pass
            else:
                cols = max(len(r) for r in all_row_cells)
                n_header = len(header_cells_rows)
                table = doc.add_table(rows=len(all_row_cells), cols=cols)
                try: table.style = 'Table Grid'
                except: pass
                for r_idx, cells in enumerate(all_row_cells):
                    is_header_row = r_idx < n_header
                    for c_idx, cell in enumerate(cells):
                        if c_idx >= cols:
                            continue
                        cell_p = table.rows[r_idx].cells[c_idx].paragraphs[0]
                        cell_p.paragraph_format.space_after = Pt(2)
                        cell_p.paragraph_format.space_before = Pt(2)
                        
                        # Tables inside callouts are tricky because cells can't easily have 
                        # half-borders and shading without messing up the table's own.
                        # Usually, if a table is inside a callout, we just let it be.
                        # But we could at least indent the table? 
                        # python-docx doesn't support table indentation easily via this API.
                        
                        if is_header_row:
                            tc = table.rows[r_idx].cells[c_idx]._tc
                            tcPr = tc.get_or_add_tcPr()
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            shd.set(qn('w:fill'), 'D9D9D9')
                            tcPr.append(shd)
                        
                        cell_children = cell.get("children", [])
                        self._add_runs(cell_p, cell_children)
                        if is_header_row:
                            for run in cell_p.runs:
                                run.bold = True

        else:
            for child in node.get("children", []):
                self._render_node(child, doc, list_level)

    def export(
        self,
        markdown_content: str,
        out_docx_path: str,
        out_pdf_path: str | None = None,
        assignment_title: str | None = None,
        course_name: str | None = None,
        metadata: dict | None = None,
        assignment_type: str | None = None,
    ):
        """Build the document using templates and export it to docx/pdf."""
        if metadata is None: metadata = {}

        is_exam = (assignment_type == "exam")
        
        # Load the initial base document from yaml config (or external if configured to do so)
        if hasattr(self, "template_path") and self.template_path and Path(self.template_path).exists():
            doc = Document(self.template_path)
            # When using the legacy external template we still need to strip the
            # placeholder paragraphs (indices >= 3) exactly as before.
            for i, p in reversed(list(enumerate(doc.paragraphs))):
                if i >= 3:
                    self._delete_paragraph(p)

            if doc.tables and not is_exam:
                if len(doc.tables) > 1:
                    t1 = doc.tables[1]
                    t1._element.getparent().remove(t1._element)
                table = doc.tables[0]
                try:
                    tr = table.rows[2]._tr
                    tr.getparent().remove(tr)
                except Exception:
                    pass
                try:
                    table.rows[1].cells[2].text = "Valor:"
                except Exception:
                    pass
        else:
            doc = self._build_base_document(is_exam, metadata)

        self.out_dir = os.path.dirname(out_docx_path)
        self.total_points = 0.0

        title_p = doc.add_paragraph()
        title_p.style = "Normal"
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(assignment_title.upper())
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(12)
        
        doc.add_paragraph() # spacing

        markdown_content = re.sub(r'^---\n.*?\n---\n', '', markdown_content, flags=re.DOTALL)

        m = mistune.create_markdown(renderer=None, plugins=["table"])
        ast = m(markdown_content)
        
        for node in ast:
            self._render_node(node, doc)
            
        curso_name = metadata.get("course")
        if not curso_name:
            if "engenharia de software" in course_name.lower():
                curso_name = "Engenharia de Computação"
            else:
                curso_name = course_name
            
        if doc.tables:
            table = doc.tables[0]
            visited_cells = set()
            for row in table.rows:
                for cell in row.cells:
                    if cell in visited_cells:
                        continue
                    visited_cells.add(cell)
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if "Curso:" in text and curso_name and f"Curso: {curso_name}" not in text:
                            paragraph.text = text.replace("Curso:", f"Curso: {curso_name}")
                        elif "Disciplina:" in text and course_name and f"Disciplina: {course_name}" not in text:
                            paragraph.text = text.replace("Disciplina:", f"Disciplina: {course_name}")
                        elif "Professor(a):" in text and "Gabriel" not in text:
                            paragraph.text = text.replace("Professor(a):", "Professor(a): Gabriel Soares Baptista")
                        elif "Valor:" in text:
                            fmt_val = str(self.total_points).replace('.', ',')
                            if fmt_val.endswith(',0'):
                                fmt_val = fmt_val[:-2] # 2,0 -> 2
                            paragraph.text = text.replace("Valor:", f"Valor: {fmt_val}")
                            
        doc.save(out_docx_path)
        
        try:
            cmd = f'libreoffice --headless --convert-to pdf "{out_docx_path}" --outdir "{self.out_dir}" > /dev/null 2>&1'
            os.system(cmd)
        except Exception as e:
            print("Failed to convert to PDF via libreoffice", e)
